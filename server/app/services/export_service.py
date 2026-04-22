import io
import os
import sys
from typing import Dict, Any, List
from jinja2 import Environment, FileSystemLoader
import logging
from app.core.config import settings

logger = logging.getLogger(__name__)

# --- Windows GTK+ Helper ---
if sys.platform == "win32":
    # Common GTK3-Runtime locations
    gtk_paths = [
        os.path.join(os.environ.get("ProgramFiles", "C:\\Program Files"), "GTK3-Runtime Win64", "bin"),
        os.path.join(os.environ.get("LOCALAPPDATA", ""), "GTK3-Runtime Win64", "bin"),
        "C:\\msys64\\mingw64\\bin",
    ]
    for path in gtk_paths:
        if os.path.exists(path) and path not in os.environ["PATH"]:
            os.environ["PATH"] = path + os.pathsep + os.environ["PATH"]
            break
# ---------------------------

# Setup Jinja2 environment
template_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
env = Environment(loader=FileSystemLoader(template_dir))

def format_skills(skills_data: List[Any]) -> List[str]:
    """Converts various skill formats into a clean list of strings."""
    formatted = []
    for skill in skills_data:
        if isinstance(skill, dict):
            formatted.append(skill.get("name", ""))
        elif isinstance(skill, str):
            formatted.append(skill)
        elif hasattr(skill, "name"):
            formatted.append(skill.name)
    return [s for s in formatted if s]

def format_experience(experience_data: List[Any]) -> List[Dict[str, Any]]:
    """Normalizes experience data for the template."""
    formatted = []
    for exp in experience_data:
        if isinstance(exp, str):
            # Fallback if it's just a string
            formatted.append({
                "role": "Professional Experience",
                "company": exp,
                "duration": "",
                "location": "",
                "bullets": []
            })
            continue
            
        role = exp.get("job_title") or exp.get("role") or exp.get("title", "Experience")
        company = exp.get("company") or exp.get("organization", "")
        duration = exp.get("duration") or exp.get("dates") or f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}"
        location = exp.get("location", "")
        
        # Handle bullets (could be 'description' as string or 'bullets'/'points' as list)
        bullets = exp.get("bullets") or exp.get("points") or []
        if not bullets and exp.get("description"):
            desc = exp.get("description")
            if isinstance(desc, list):
                bullets = desc
            else:
                # Split by newlines if it's a block of text
                bullets = [b.strip() for b in desc.split("\n") if b.strip()]
        
        formatted.append({
            "role": role,
            "company": company,
            "duration": duration.strip(" -"),
            "location": location,
            "bullets": bullets
        })
    return formatted

def format_education(education_data: List[Any]) -> List[Dict[str, Any]]:
    """Normalizes education data for the template."""
    formatted = []
    for edu in education_data:
        if isinstance(edu, str):
            formatted.append({
                "degree": edu,
                "school": "",
                "duration": "",
                "location": ""
            })
            continue
            
        degree = edu.get("degree") or edu.get("qualification", "Education")
        school = edu.get("institution") or edu.get("university") or edu.get("school", "")
        duration = edu.get("duration") or edu.get("dates") or f"{edu.get('start_year', '')} - {edu.get('end_year', '')}"
        location = edu.get("location", "")
        
        formatted.append({
            "degree": degree,
            "school": school,
            "duration": duration.strip(" -"),
            "location": location
        })
    return formatted

def format_projects(projects_data: List[Any]) -> List[Dict[str, Any]]:
    """Normalizes projects data for the template."""
    formatted = []
    for proj in projects_data:
        if isinstance(proj, str):
            formatted.append({
                "name": proj,
                "date": "",
                "bullets": []
            })
            continue
            
        name = proj.get("title") or proj.get("name", "Project")
        date = proj.get("date") or proj.get("duration", "")
        
        bullets = proj.get("bullets") or []
        if not bullets and proj.get("description"):
            desc = proj.get("description")
            if isinstance(desc, list):
                bullets = desc
            else:
                bullets = [b.strip() for b in desc.split("\n") if b.strip()]
                
        formatted.append({
            "name": name,
            "date": date,
            "bullets": bullets
        })
    return formatted

def generate_resume_html(resume_data: Dict[str, Any]) -> str:
    """Loads the template and injects formatted resume data."""
    template = env.get_template("resume_template.html")
    
    # Prepare data for injection
    contact_str = resume_data.get("contact", "")
    parts = [p.strip() for p in contact_str.split("|")]
    
    # Prepare tagline
    experience = resume_data.get("experience", [])
    tagline = ""
    if experience and isinstance(experience, list):
        first_exp = experience[0]
        if isinstance(first_exp, dict):
            tagline = first_exp.get("job_title") or first_exp.get("role") or first_exp.get("title")
    
    context = {
        "name": resume_data.get("name", "Name Not Provided"),
        "tagline": tagline or "Professional Resume",
        "phone": resume_data.get("phone") or (parts[0] if len(parts) > 0 else ""),
        "email": resume_data.get("email") or (parts[1] if len(parts) > 1 else ""),
        "location": resume_data.get("location") or (parts[2] if len(parts) > 2 else ""),
        "contact": contact_str,
        "summary": resume_data.get("summary", ""),
        "skills": format_skills(resume_data.get("skills", [])),
        "experience": format_experience(resume_data.get("experience", [])),
        "education": format_education(resume_data.get("education", [])),
        "projects": format_projects(resume_data.get("projects", []))
    }
    
    return template.render(context)

def generate_resume_pdf(resume_data: Dict[str, Any]) -> bytes:
    """Generates a professional PDF using WeasyPrint with a fallback to xhtml2pdf."""
    html_content = generate_resume_html(resume_data)
    
    # Try WeasyPrint first (better quality)
    try:
        from weasyprint import HTML
        return HTML(string=html_content).write_pdf()
    except Exception as e:
        logger.error(f"WeasyPrint failed: {e}. Trying fallback...")
        
        # Fallback to xhtml2pdf (pure python, works on Windows without GTK)
        try:
            from xhtml2pdf import pisa
            result = io.BytesIO()
            pisa_status = pisa.CreatePDF(io.StringIO(html_content), dest=result)
            if not pisa_status.err:
                return result.getvalue()
            else:
                logger.error(f"xhtml2pdf error: {pisa_status.err}")
        except Exception as e2:
            logger.error(f"Fallback xhtml2pdf failed: {e2}")
            
        # If both fail, raise the original WeasyPrint error with helpful advice
        if "cannot load library" in str(e).lower() or "gobject" in str(e).lower():
            raise RuntimeError(
                "PDF Export failed: WeasyPrint dependencies (GTK+) are missing on your Windows system. "
                "You can still export as DOCX, or install GTK+ to enable high-quality PDF export."
            ) from e
        raise e

def generate_resume_docx(resume_data: Dict[str, Any]) -> bytes:
    """Generates a professional DOCX using python-docx as a fallback/alternative."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Set styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(11)

    # Header
    header = doc.add_heading(resume_data.get('name', 'Name Not Provided'), 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_str = resume_data.get("contact") or f"{resume_data.get('email', '')} | {resume_data.get('phone', '')}".strip(" |")
    contact = doc.add_paragraph(contact_str)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    if resume_data.get("summary"):
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        doc.add_paragraph(resume_data.get("summary"))

    # Skills
    doc.add_heading('TECHNICAL SKILLS', level=1)
    skills = format_skills(resume_data.get("skills", []))
    doc.add_paragraph(" • ".join(skills))

    # Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    for exp in format_experience(resume_data.get('experience', [])):
        p = doc.add_paragraph()
        run = p.add_run(f"{exp['role']} | {exp['company']}")
        run.bold = True
        p.add_run(f"\t{exp['duration']}")
        
        for bullet in exp['bullets']:
            doc.add_paragraph(bullet, style='List Bullet')

    # Projects
    if resume_data.get("projects"):
        doc.add_heading('PROJECTS', level=1)
        for proj in format_projects(resume_data.get('projects', [])):
            p = doc.add_paragraph()
            run = p.add_run(proj['name'])
            run.bold = True
            p.add_run(f"\t{proj['date']}")
            
            for bullet in proj['bullets']:
                doc.add_paragraph(bullet, style='List Bullet')

    # Education
    doc.add_heading('EDUCATION', level=1)
    for edu in format_education(resume_data.get('education', [])):
        p = doc.add_paragraph()
        run = p.add_run(edu['degree'])
        run.bold = True
        p.add_run(f"\t{edu['duration']}")
        doc.add_paragraph(edu['school'])

    # Save to bytes
    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

